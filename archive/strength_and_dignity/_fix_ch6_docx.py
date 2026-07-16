#!/usr/bin/env python3
"""
Sync StrengthAndDignity_Chapter6.docx with the HTML version.

Two edits that were applied to chapter-06.html but missed the docx
source — meaning every regenerated reader PDF / EPUB / Lulu interior
still carried the old text:

  1. Delete the duplicate Psalm 119:105 quote + the "lamp does not
     illuminate the whole road / next step" paragraph that follows it.
     (The same quote already lives in Ch 4. The lamp/next-step gloss
     is also on the cliche-avoid list.)

  2. Trim "— and it should stop you in your tracks" from the
     "Living and Active" lead-in. The HTML version ends that sentence
     with a colon after "actually is".
"""

from pathlib import Path
from docx import Document

SRC = Path(__file__).parent / "StrengthAndDignity_Chapter6.docx"
assert SRC.exists(), SRC

doc = Document(str(SRC))

# ---- 1. Drop the three duplicate paragraphs ------------------------------
DROP_PREFIXES = (
    "“Your word is a lamp",          # the quote
    "— Psalm 119:105",                # the citation
    "A lamp does not illuminate",     # the AI-tic lamp paragraph
)

removed = []
for para in list(doc.paragraphs):
    text = para.text.strip()
    if text.startswith(DROP_PREFIXES):
        para._element.getparent().remove(para._element)
        removed.append(text[:80])

print("Removed paragraphs:")
for r in removed:
    print(f"  - {r}")
assert len(removed) == 3, f"Expected to remove 3 paragraphs, removed {len(removed)}"

# ---- 2. Trim the "stop you in your tracks" tail --------------------------
TARGET_OLD = "There is another passage that describes what the Bible actually is — and it should stop you in your tracks:"
TARGET_NEW = "There is another passage that describes what the Bible actually is:"

found = False
for para in doc.paragraphs:
    if para.text.strip() == TARGET_OLD:
        # Replace by rewriting the run text. Use the simplest approach:
        # collapse to a single run with the new text, preserving the
        # paragraph's run formatting from the first run.
        first_run = para.runs[0] if para.runs else None
        # Clear all runs
        for r in list(para.runs):
            r.text = ""
        if first_run is not None:
            first_run.text = TARGET_NEW
        else:
            para.add_run(TARGET_NEW)
        found = True
        print(f"\nRewrote paragraph:\n  OLD: {TARGET_OLD}\n  NEW: {TARGET_NEW}")
        break

assert found, "Did not find the 'stop you in your tracks' paragraph to rewrite."

doc.save(str(SRC))
print(f"\nSaved {SRC}")

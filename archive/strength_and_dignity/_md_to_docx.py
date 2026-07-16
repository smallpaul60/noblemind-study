#!/usr/bin/env python3
"""
Convert StrengthAndDignity_Chapter9.md (the new Chapter 9 source) into
StrengthAndDignity_Chapter9.docx in the format expected by generate_pdf.py /
generate_lulu_interior.py.

Format conventions used by the docx parser in those generators:

  - All-bold paragraph (short text)              -> h2 section heading
  - All-italic paragraph starting with " or U+201C  -> scripture quote
  - Paragraph starting with em-dash              -> citation line for the
    scripture quote immediately above
  - Bold-italic paragraph                        -> principle pull-quote
    (rendered as <div class="principle-box">)
  - Paragraph beginning with U+2022 (•) bullet   -> list item
  - "For Further Study" (bold-italic or bold)    -> end-of-chapter section header

The first two paragraphs of every chapter are bold (CHAPTER N + title) and the
parser skips them; we emit them in the same form for consistency.
"""

import re
import sys
from pathlib import Path

from docx import Document

SRC = Path(__file__).parent / "StrengthAndDignity_Chapter9.md"
DST = Path(__file__).parent / "StrengthAndDignity_Chapter9.docx"


def parse_markdown(text: str):
    """Yield (kind, payload) tuples for each non-blank logical paragraph."""
    blocks = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue
        blocks.append(line)
    return blocks


def classify(line: str):
    """
    Classify a markdown paragraph into one of:
      ('bold',       text)            — wrapped in **...**
      ('bold_italic', text)           — wrapped in ***...***
      ('italic_quote', text)          — wrapped in *"..."*  (scripture)
      ('italic_plain', text)          — wrapped in *...*    (no quote marks)
      ('citation',   text)            — starts with em-dash
      ('bullet',     text)            — starts with •  (the source uses
                                        **•  **text — a bold bullet then plain)
      ('plain',      text)            — anything else
    """
    # Bold-italic: ***...***
    m = re.fullmatch(r"\*\*\*(.+?)\*\*\*", line)
    if m:
        return ("bold_italic", m.group(1))

    # Bold: **...**
    m = re.fullmatch(r"\*\*(.+?)\*\*", line)
    if m:
        return ("bold", m.group(1))

    # Italic with quote (scripture). Match *"..."*
    m = re.fullmatch(r"\*[“\"](.+?)[”\"]\*", line)
    if m:
        return ("italic_quote", m.group(1))

    # Italic plain (no quotes): *...*
    m = re.fullmatch(r"\*(.+?)\*", line)
    if m:
        return ("italic_plain", m.group(1))

    # Citation: starts with em-dash
    if line.startswith("—") or line.startswith("—"):
        # strip the leading em-dash + space if any
        return ("citation", line)

    # Bullet: source uses "**•  **TEXT" — strip optional ** wrapping the bullet,
    # then the line begins with • + space + content.
    bullet_match = re.match(r"^(?:\*\*)?[••]\s*(?:\*\*)?\s*(.+)$", line)
    if bullet_match and (line.startswith("**•") or line.startswith("•") or line.startswith("•")):
        return ("bullet", bullet_match.group(1))

    return ("plain", line)


def add_run(paragraph, text: str, *, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    return run


def emit_plain_with_inline(doc, text: str):
    """
    Add a plain paragraph, parsing inline *italic* spans and **bold** spans.
    Used for body prose so that <em>...</em> emphasis survives into the docx.
    """
    p = doc.add_paragraph()
    # Inline tokenizer: split on ** or * markers (greedy, non-overlapping).
    # We support ***text*** (both), **text** (bold), *text* (italic).
    # Order matters in the regex alternatives.
    pattern = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)", re.DOTALL)
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_run(p, text[pos:m.start()])
        token = m.group(0)
        if token.startswith("***"):
            add_run(p, token[3:-3], bold=True, italic=True)
        elif token.startswith("**"):
            add_run(p, token[2:-2], bold=True)
        else:
            add_run(p, token[1:-1], italic=True)
        pos = m.end()
    if pos < len(text):
        add_run(p, text[pos:])


def main():
    raw = SRC.read_text(encoding="utf-8")
    blocks = parse_markdown(raw)

    doc = Document()

    for line in blocks:
        kind, payload = classify(line)

        if kind == "bold":
            p = doc.add_paragraph()
            add_run(p, payload, bold=True)

        elif kind == "bold_italic":
            p = doc.add_paragraph()
            add_run(p, payload, bold=True, italic=True)

        elif kind == "italic_quote":
            # Wrap with curly quotes to ensure parser recognizes scripture.
            p = doc.add_paragraph()
            add_run(p, "“" + payload + "”", italic=True)

        elif kind == "italic_plain":
            # Italic without quotes — body emphasis (e.g., the closing line
            # before For Further Study). Parser falls through to format_inline
            # and preserves italic via run.italic.
            p = doc.add_paragraph()
            add_run(p, payload, italic=True)

        elif kind == "citation":
            p = doc.add_paragraph()
            add_run(p, payload)  # plain run; parser sees em-dash prefix

        elif kind == "bullet":
            # Emit as plain paragraph beginning with the bullet character —
            # parser collects these into a <ul>.
            p = doc.add_paragraph()
            add_run(p, "•  ")
            # Allow inline emphasis inside the bullet text just in case.
            inline_pattern = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)", re.DOTALL)
            pos = 0
            for m in inline_pattern.finditer(payload):
                if m.start() > pos:
                    add_run(p, payload[pos:m.start()])
                tok = m.group(0)
                if tok.startswith("***"):
                    add_run(p, tok[3:-3], bold=True, italic=True)
                elif tok.startswith("**"):
                    add_run(p, tok[2:-2], bold=True)
                else:
                    add_run(p, tok[1:-1], italic=True)
                pos = m.end()
            if pos < len(payload):
                add_run(p, payload[pos:])

        else:  # plain
            emit_plain_with_inline(doc, payload)

    doc.save(str(DST))
    print(f"wrote {DST}")


if __name__ == "__main__":
    main()
